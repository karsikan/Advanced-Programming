import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    print(f"Generating {docx_path} from {md_path}...")
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('![') and '](' in line:
            # Simple image handling
            try:
                img_name = line.split('](')[1].split(')')[0].strip('/')
                img_path = os.path.join(os.path.dirname(md_path), img_name)
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(5.5))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph(f"[Image Missing: {img_name}]")
            except Exception as e:
                doc.add_paragraph(f"[Error loading image: {str(e)}]")
        elif line.startswith('|') and '|' in line:
            # Simple table detection
            if '---' in line: continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not hasattr(doc, '_last_table_line'):
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                row = table.rows[0]
                for i, c in enumerate(cells):
                    row.cells[i].text = c
            else:
                table = doc.tables[-1]
                row = table.add_row()
                for i, c in enumerate(cells):
                    if i < len(row.cells):
                        row.cells[i].text = c
        else:
            p = doc.add_paragraph(line)
            if '**' in line: # Bold detection
                pass # Simple text for now
            
    doc.save(docx_path)
    print(f"Successfully saved {docx_path}")

artifact_dir = r"c:\Users\hp\.gemini\antigravity\brain\333b277f-3430-4d33-8c00-0749fb3c481f"
reports = [
    ("UML_Design_Documentation.md", "UML_Design_Report_Final.docx"),
    ("Technical_Implementation_Report.md", "Technical_Implementation_Report_Final.docx"),
    ("Test_Plan_and_Report.md", "Test_Plan_and_Report_Final.docx")
]

for md, docx in reports:
    markdown_to_docx(os.path.join(artifact_dir, md), os.path.join(artifact_dir, docx))
